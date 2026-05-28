"""Word document generator for math problems."""

from __future__ import annotations

from datetime import date
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.shared import Pt, Cm

from math_learning.core.generator import Problem

# Layout constants
COLS = 4
FONT_SIZE = Pt(14)
TITLE_FONT_SIZE = Pt(22)
ROW_SPACING = Pt(28)


def generate_word(problems: list[Problem], title: str = "口算练习题") -> BytesIO:
    """Generate a Word document with math problems in a grid layout.

    Args:
        problems: List of Problem instances to include.
        title: Document title.

    Returns:
        BytesIO containing the .docx file bytes.
    """
    doc = Document()

    # Page setup - A4
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.font.size = TITLE_FONT_SIZE
    title_run.bold = True
    title_para.space_after = Pt(6)

    # Subtitle: name / class / date line
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_para.add_run("姓名：__________    班级：__________    日期：__________")
    info_run.font.size = Pt(11)
    info_para.space_after = Pt(16)

    # Problems in table layout (4 columns)
    rows_needed = (len(problems) + COLS - 1) // COLS
    table = doc.add_table(rows=rows_needed, cols=COLS)
    table.autofit = True

    for idx, problem in enumerate(problems):
        row_idx = idx // COLS
        col_idx = idx % COLS
        cell = table.cell(row_idx, col_idx)
        para = cell.paragraphs[0]
        run = para.add_run(problem.expression)
        run.font.size = FONT_SIZE
        run.font.name = "Consolas"
        para.space_after = ROW_SPACING

    # Footer: generation info
    doc.add_paragraph()  # spacer
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer_para.add_run(
        f"共 {len(problems)} 题  |  生成日期：{date.today().isoformat()}"
    )
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = None  # default

    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
